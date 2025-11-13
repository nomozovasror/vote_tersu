from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from typing import List


def set_cell_border(cell, **kwargs):
    """Set border for table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)


def generate_results_word(event_name: str, results: List[dict], total_participants: int) -> BytesIO:
    """Generate Word document with voting results"""
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run(f"{event_name} - Ovoz berish natijalari")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing
    doc.add_paragraph()

    # Create table with 8 columns
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    # Set column widths (approximate)
    widths = [Inches(0.5), Inches(2.0), Inches(2.0), Inches(1.0),
              Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.0)]

    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # Header row
    header_cells = table.rows[0].cells
    headers = ['T/r', 'Lavozim', 'Nomzodlar', 'Ovoz berishda qatnashganlar soni',
               'Rozi', 'Qarshi', 'Betaraf', 'Natija']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text

        # Format header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set cell shading (light gray)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading_elm)

        # Add borders
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "000000"},
            bottom={"sz": 12, "val": "single", "color": "000000"},
            left={"sz": 12, "val": "single", "color": "000000"},
            right={"sz": 12, "val": "single", "color": "000000"}
        )

    # Add data rows
    for result in results:
        row = table.add_row()
        cells = row.cells

        # T/r (row number)
        cells[0].text = str(result.get('row_number', ''))

        # Lavozim (position)
        cells[1].text = result.get('which_position', '')

        # Nomzodlar (candidate name)
        cells[2].text = result.get('full_name', '')

        # Ovoz berishda qatnashganlar soni (yes + no + neutral for this candidate)
        yes_votes = result.get('yes_votes', 0)
        no_votes = result.get('no_votes', 0)
        neutral_votes = result.get('neutral_votes', 0)
        candidate_participants = yes_votes + no_votes + neutral_votes
        cells[3].text = str(candidate_participants)

        # Rozi (yes votes with percentage)
        yes_percent = result.get('yes_percent', 0)
        cells[4].text = f"{yes_votes}\n({yes_percent}%)"

        # Qarshi (no votes with percentage)
        no_percent = result.get('no_percent', 0)
        cells[5].text = f"{no_votes}\n({no_percent}%)"

        # Betaraf (neutral votes with percentage)
        neutral_percent = result.get('neutral_percent', 0)
        cells[6].text = f"{neutral_votes}\n({neutral_percent}%)"

        # Natija (result)
        cells[7].text = result.get('result', '')

        # Format cells
        for i, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                # Center align row number and numbers
                if i in [0, 3, 4, 5, 6, 7]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add borders
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"}
            )

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
